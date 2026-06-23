"""
ingest.py
=========

"""

import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
import time
from html.parser import HTMLParser
from pathlib import Path
from typing import Dict, List, Optional, Sequence

import requests

try:
    from bs4 import BeautifulSoup
except ModuleNotFoundError:
    BeautifulSoup = None

from config import (
    BM25_INDEX_PATH,
    CHUNK_MIN_WORDS,
    CHUNK_OVERLAP,
    CHUNK_WORDS,
    DATA_DIR,
    DOCSTORE_PATH,
    EMBEDDING_DIM,
    FAISS_INDEX_PATH,
    INGEST_ALLOWED_ROOT,
)
from rag import embed_texts, load_docstore, load_index, save_index

logger = logging.getLogger(__name__)

try:
    import fitz  # type: ignore
except ModuleNotFoundError:  # pragma: no cover - optional dependency
    fitz = None

_FITZ_EMPTY_FILE_ERROR = getattr(fitz, "EmptyFileError", ValueError)


# ── Directory helpers ──────────────────────────────────────────────────────


def _ensure_dirs() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    FAISS_INDEX_PATH.parent.mkdir(parents=True, exist_ok=True)


# ── Text utilities ─────────────────────────────────────────────────────────


def clean_text(text: str) -> str:
    """Strip null bytes and collapse whitespace."""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(
    text: str,
    chunk_words: int = CHUNK_WORDS,
    overlap: int = CHUNK_OVERLAP,
) -> List[str]:
    """
    Split *text* into overlapping word-count chunks.
    """
    words = clean_text(text).split()
    if not words:
        return []

    chunks: List[str] = []
    start = 0
    max_merge_words = chunk_words * 2

    while start < len(words):
        end = min(start + chunk_words, len(words))
        chunk = " ".join(words[start:end]).strip()
        if chunk:
            word_count = len(chunk.split())
            if word_count < CHUNK_MIN_WORDS and chunks:
                prev_word_count = len(chunks[-1].split())
                if prev_word_count + word_count <= max_merge_words:
                    chunks[-1] = f"{chunks[-1]} {chunk}".strip()
                else:
                    chunks.append(chunk)
            else:
                chunks.append(chunk)
        if end >= len(words):
            break
        start = max(0, end - overlap)

    return chunks


def chunk_text_semantic(
    text: str,
    chunk_words: int = CHUNK_WORDS,
    overlap: int = CHUNK_OVERLAP,
) -> List[str]:
    """Split text into sentence-bounded chunks with light overlap."""
    try:
        cleaned = clean_text(text)
        sentences = [
            part.strip()
            for part in re.split(r"(?<=[.!?])\s+|\n{2,}", cleaned)
            if part and part.strip()
        ]
        if not sentences:
            return []

        sentence_words = [len(sentence.split()) for sentence in sentences]
        chunks: List[str] = []
        start = 0
        max_words = max(chunk_words, CHUNK_MIN_WORDS)

        while start < len(sentences):
            end = start
            total = 0
            window: List[str] = []
            while end < len(sentences):
                sentence = sentences[end]
                count = sentence_words[end]
                if window and total + count > max_words:
                    break
                window.append(sentence)
                total += count
                end += 1
                if total >= max_words:
                    break

            chunk = " ".join(window).strip()
            if chunk:
                if len(chunk.split()) < CHUNK_MIN_WORDS and chunks:
                    chunks[-1] = f"{chunks[-1]} {chunk}".strip()
                else:
                    chunks.append(chunk)

            if end == start:
                end = start + 1

            if end >= len(sentences):
                break

            overlap_words = 0
            new_start = end
            while new_start > start and overlap_words < overlap:
                new_start -= 1
                overlap_words += sentence_words[new_start]
            start = max(new_start, start + 1)

        return chunks
    except ValueError as exc:
        logger.warning("Semantic chunking failed with ValueError: %s", exc)
        raise
    except OSError as exc:
        logger.warning("Semantic chunking failed with OSError: %s", exc)
        raise
    except _FITZ_EMPTY_FILE_ERROR as exc:
        logger.warning("Semantic chunking failed with EmptyFileError: %s", exc)
        raise


def _normalise_source(source: str) -> str:
    """Normalise a source path/URL for deduplication comparisons."""
    source = (source or "").strip()
    if "://" in source:
        return source.rstrip("/")
    if source.startswith("manual_text"):
        return source.rstrip("/")
    return str(Path(source).as_posix()).rstrip("/")


def _path_source(path: Path) -> str:
    """Return a resolved path as a POSIX source string."""
    return str(path.expanduser().resolve().as_posix())


def _source_already_ingested(source: str) -> bool:
    """Return True if *source* is already in the docstore."""
    if not source:
        return False
    docs = load_docstore()
    normalised = _normalise_source(source)
    return any(_normalise_source(d.get("source", "")) == normalised for d in docs)


# ── Core append ───────────────────────────────────────────────────────────


def _append_documents(
    chunks: Sequence[str],
    source: str,
    doc_type: str,
    original_filename: Optional[str] = None,
) -> int:
    """
    Embed *chunks* and append them to the FAISS index + docstore.
    Returns number of chunks actually added.
    """
    if not chunks:
        return 0

    try:
        _ensure_dirs()
        normalized_source = _normalise_source(source)
        index = load_index()
        docs = load_docstore()

        vectors = embed_texts(chunks)
        if vectors.size == 0:
            return 0

        if index.ntotal == 0 and vectors.shape[1] != EMBEDDING_DIM:
            raise ValueError(
                f"Embedding dimension mismatch: expected {EMBEDDING_DIM}, "
                f"got {vectors.shape[1]}"
            )

        index.add(vectors)

        start_id = len(docs) + 1
        for i, chunk in enumerate(chunks, start=start_id):
            entry: Dict = {
                "source": normalized_source,
                "doc_type": doc_type,
                "chunk_id": str(i),
                "text": chunk,
                "ingested_at": time.time(),
            }
            if original_filename:
                entry["original_filename"] = original_filename
            docs.append(entry)

        save_index(index, docs)
        return len(chunks)
    except ValueError as exc:
        logger.warning("Appending documents failed with ValueError: %s", exc)
        raise
    except OSError as exc:
        logger.warning("Appending documents failed with OSError: %s", exc)
        raise
    except json.JSONDecodeError as exc:
        logger.warning("Appending documents failed with JSONDecodeError: %s", exc)
        raise
    except _FITZ_EMPTY_FILE_ERROR as exc:
        logger.warning("Appending documents failed with EmptyFileError: %s", exc)
        raise


# ── HTML extractor ─────────────────────────────────────────────────────────


class _VisibleTextExtractor(HTMLParser):
    _BLOCK_TAGS = {"h1", "h2", "h3", "h4", "h5", "p", "li", "td", "th", "br", "div"}
    _SKIP_TAGS = {"script", "style", "noscript", "header", "footer", "nav"}

    def __init__(self) -> None:
        super().__init__()
        self.parts: List[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP_TAGS:
            self._skip_depth += 1
        elif (
            tag in self._BLOCK_TAGS and self.parts and not self.parts[-1].endswith("\n")
        ):
            self.parts.append("\n")

    def handle_endtag(self, tag):
        if tag in self._SKIP_TAGS and self._skip_depth > 0:
            self._skip_depth -= 1
        elif (
            tag in self._BLOCK_TAGS and self.parts and not self.parts[-1].endswith("\n")
        ):
            self.parts.append("\n")

    def handle_data(self, data):
        if self._skip_depth == 0:
            text = data.strip()
            if text:
                self.parts.append(text + " ")


def _extract_visible_text(html: str) -> str:
    if BeautifulSoup is not None:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
            tag.decompose()
        parts: List[str] = []
        for element in soup.find_all(
            ["h1", "h2", "h3", "h4", "h5", "p", "li", "td", "th"]
        ):
            text = element.get_text(" ", strip=True)
            if text:
                parts.append(text)
        return " ".join(parts)

    parser = _VisibleTextExtractor()
    parser.feed(html)
    parser.close()
    return clean_text("".join(parser.parts))


# ── Legacy .doc extractor (Word 97-2003) ──────────────────────────────────


def _extract_doc_text(file_path: str) -> str:
    """
    Extract text from a legacy .doc file.

    Priority:
      1. LibreOffice headless  → convert .doc to .docx, then read with python-docx
      2. antiword              → lightweight CLI tool, returns plain text
      3. PyMuPDF (fitz)        → works on some .doc files as a last resort

    Raises ValueError if no method succeeds or no text is extracted.
    """
    path = Path(file_path).expanduser().resolve()

    # ── Option 1: LibreOffice ──────────────────────────────────────────────
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run(
                    [
                        soffice,
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        tmpdir,
                        str(path),
                    ],
                    capture_output=True,
                    timeout=60,
                )
                converted_name = path.stem + ".docx"
                converted_path = Path(tmpdir) / converted_name
                if converted_path.exists():
                    try:
                        from docx import Document as DocxDocument  # type: ignore

                        doc = DocxDocument(str(converted_path))
                        paragraphs = [
                            p.text.strip() for p in doc.paragraphs if p.text.strip()
                        ]
                        text = clean_text("\n".join(paragraphs))
                        if text:
                            logger.info("Extracted .doc via LibreOffice: %s", path.name)
                            return text
                    except Exception as exc:
                        logger.debug("LibreOffice-converted docx read failed: %s", exc)
        except Exception as exc:
            logger.debug("LibreOffice .doc conversion failed: %s", exc)

    # ── Option 2: antiword ─────────────────────────────────────────────────
    antiword = shutil.which("antiword")
    if antiword:
        try:
            result = subprocess.run(
                [antiword, str(path)],
                capture_output=True,
                text=True,
                timeout=30,
            )
            text = clean_text(result.stdout)
            if text:
                logger.info("Extracted .doc via antiword: %s", path.name)
                return text
        except Exception as exc:
            logger.debug("antiword extraction failed: %s", exc)

    # ── Option 3: PyMuPDF fallback ─────────────────────────────────────────
    if fitz is not None:
        try:
            doc = fitz.open(str(path))
            pages = [
                page.get_text("text") for page in doc if page.get_text("text").strip()
            ]
            text = clean_text("\n".join(pages))
            if text:
                logger.info("Extracted .doc via PyMuPDF fallback: %s", path.name)
                return text
        except Exception as exc:
            logger.debug("PyMuPDF .doc fallback failed: %s", exc)

    raise ValueError(
        f"Could not extract text from '{path.name}'. "
        "Install LibreOffice (soffice) or antiword on the server to support .doc files."
    )


# ── Public ingest functions ────────────────────────────────────────────────


def ingest_pdf(
    file_path: str,
    force: bool = False,
    original_filename: Optional[str] = None,
) -> int:
    """Extract text from a PDF and add it to the knowledge base."""
    path = Path(file_path).expanduser().resolve()
    # Security Fix: The old code lacked path containment validation, allowing arbitrary local file read.
    if not path.is_relative_to(INGEST_ALLOWED_ROOT):
        raise ValueError(f"Path '{path}' is not within the allowed ingest root directory.")
    if not path.exists():
        raise FileNotFoundError(f"PDF not found: {path}")
    if path.suffix.lower() != ".pdf":
        raise ValueError(f"Expected a .pdf file, got: {path.suffix}")

    if fitz is None:
        raise RuntimeError("PyMuPDF is not installed. Run: pip install PyMuPDF")

    source = _path_source(path)
    dedup_source = original_filename if original_filename else source
    if not force and _source_already_ingested(dedup_source):
        return 0

    pages: List[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            page_text = page.get_text("text")
            if page_text and page_text.strip():
                pages.append(page_text)

    if not pages:
        raise ValueError(
            "The PDF contains no extractable text. "
            "It may be a scanned/image-only PDF."
        )

    text = clean_text("\n".join(pages))
    chunks = chunk_text_semantic(text)
    return _append_documents(
        chunks,
        source=dedup_source,
        doc_type="pdf",
        original_filename=original_filename,
    )


def ingest_txt(
    file_path: str,
    force: bool = False,
    original_filename: Optional[str] = None,
) -> int:
    """Ingest a plain .txt file."""
    path = Path(file_path).expanduser().resolve()
    # Security Fix: The old code lacked path containment validation, allowing arbitrary local file read.
    if not path.is_relative_to(INGEST_ALLOWED_ROOT):
        raise ValueError(f"Path '{path}' is not within the allowed ingest root directory.")
    if not path.exists():
        raise FileNotFoundError(f"Text file not found: {path}")

    source = _path_source(path)
    dedup_source = original_filename if original_filename else source
    if not force and _source_already_ingested(dedup_source):
        return 0

    text = path.read_text(encoding="utf-8", errors="replace")
    chunks = chunk_text_semantic(text)
    return _append_documents(
        chunks,
        source=dedup_source,
        doc_type="txt",
        original_filename=original_filename,
    )


def ingest_docx(
    file_path: str,
    force: bool = False,
    original_filename: Optional[str] = None,
) -> int:
    """Ingest a Microsoft Word (.docx) file."""
    path = Path(file_path).expanduser().resolve()
    # Security Fix: The old code lacked path containment validation, allowing arbitrary local file read.
    if not path.is_relative_to(INGEST_ALLOWED_ROOT):
        raise ValueError(f"Path '{path}' is not within the allowed ingest root directory.")
    if not path.exists():
        raise FileNotFoundError(f"Word file not found: {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Expected a .docx file, got: {path.suffix}")

    try:
        from docx import Document as DocxDocument  # type: ignore
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    source = _path_source(path)
    dedup_source = original_filename if original_filename else source
    if not force and _source_already_ingested(dedup_source):
        return 0

    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = clean_text("\n".join(paragraphs))
    if not text:
        raise ValueError("No extractable text found in the Word document.")

    chunks = chunk_text_semantic(text)
    return _append_documents(
        chunks,
        source=dedup_source,
        doc_type="docx",
        original_filename=original_filename,
    )


def ingest_doc(
    file_path: str,
    force: bool = False,
    original_filename: Optional[str] = None,
) -> int:
    """
    Ingest a legacy Microsoft Word (.doc) file.

    Extraction priority:
      1. LibreOffice (convert to .docx then parse with python-docx)
      2. antiword (plain-text CLI extraction)
      3. PyMuPDF fitz (last-resort fallback)

    Raises RuntimeError with a helpful message if none of the tools are available.
    """
    path = Path(file_path).expanduser().resolve()
    # Security Fix: The old code lacked path containment validation, allowing arbitrary local file read.
    if not path.is_relative_to(INGEST_ALLOWED_ROOT):
        raise ValueError(f"Path '{path}' is not within the allowed ingest root directory.")
    if not path.exists():
        raise FileNotFoundError(f"Word .doc file not found: {path}")
    if path.suffix.lower() != ".doc":
        raise ValueError(f"Expected a .doc file, got: {path.suffix}")

    source = _path_source(path)
    dedup_source = original_filename if original_filename else source
    if not force and _source_already_ingested(dedup_source):
        return 0

    text = _extract_doc_text(str(path))
    if not text:
        raise ValueError("No extractable text found in the Word .doc file.")

    chunks = chunk_text_semantic(text)
    return _append_documents(
        chunks,
        source=dedup_source,
        doc_type="doc",
        original_filename=original_filename,
    )


def _validate_url(url: str) -> None:
    # Security Fix: The old code lacked target host/IP validation, enabling SSRF attacks against internal networks.
    if os.getenv("ALLOW_INTERNAL_INGEST_URLS") == "1":
        return

    from urllib.parse import urlparse
    import socket
    import ipaddress

    parsed = urlparse(url)
    hostname = parsed.hostname
    if not hostname:
        raise ValueError("Invalid URL: missing hostname.")

    try:
        ips = socket.getaddrinfo(hostname, None)
    except Exception as e:
        raise ValueError(f"Failed to resolve hostname '{hostname}': {e}")

    for item in ips:
        ip = item[4][0]
        try:
            ip_obj = ipaddress.ip_address(ip)
            if ip_obj.is_loopback or ip_obj.is_private or ip_obj.is_link_local:
                raise ValueError(f"URL resolves to a restricted IP address: {ip}")
        except ValueError as ve:
            if "resolves to a restricted IP address" in str(ve):
                raise
            raise ValueError(f"Invalid IP address resolved: {ip}")


def ingest_url(url: str, force: bool = False) -> int:
    """Fetch a webpage, extract visible text, and add it to the knowledge base."""
    if not force and _source_already_ingested(url):
        return 0

    # Security Fix: The old code had no host/IP checks, allowing SSRF.
    _validate_url(url)

    session = requests.Session()
    session.max_redirects = 5
    response = session.get(
        url,
        allow_redirects=True,
        timeout=(20, 30),
        headers={"User-Agent": "AgniAI/1.0 (offline-chatbot)"},
    )
    response.raise_for_status()

    text = clean_text(_extract_visible_text(response.text))
    if not text:
        raise ValueError("No readable text found at the URL.")

    chunks = chunk_text_semantic(text)
    return _append_documents(chunks, source=url, doc_type="url")


def ingest_text(text: str, label: str = "manual_text") -> int:
    """
    Ingest raw text directly (e.g. pasted content).
    """
    if label == "manual_text":
        docs = load_docstore()
        existing_count = sum(
            1 for d in docs if d.get("source", "").startswith("manual_text")
        )
        if existing_count > 0:
            label = f"manual_text_{existing_count + 1}_{time.time_ns()}"

    if _source_already_ingested(label):
        return 0

    chunks = chunk_text_semantic(text)
    return _append_documents(chunks, source=label, doc_type="text")


def list_sources() -> List[Dict]:
    """Return a summary of all ingested sources."""
    docs = load_docstore()
    counts: Dict[str, Dict] = {}
    for d in docs:
        src = d.get("source", "unknown")
        if src not in counts:
            counts[src] = {
                "source": src,
                "original_filename": d.get("original_filename", ""),
                "doc_type": d.get("doc_type", "?"),
                "chunk_count": 0,
            }
        counts[src]["chunk_count"] += 1
    return list(counts.values())


def clear_index() -> None:
    """Delete indexed data and clear FAISS/docstore plus query, retrieval, response, and BM25 caches."""
    import rag

    with rag._INDEX_LOCK:
        if FAISS_INDEX_PATH.exists():
            FAISS_INDEX_PATH.unlink()
        if DOCSTORE_PATH.exists():
            DOCSTORE_PATH.unlink()
        if BM25_INDEX_PATH.exists():
            BM25_INDEX_PATH.unlink()
        rag._INDEX = None
        rag._DOCS = []
        rag._BM25 = None
        rag._DOCSTORE_CACHE = None

    rag._QUERY_EMBED_CACHE.clear()
    rag._RETRIEVAL_CACHE.clear()
    rag._RESPONSE_CACHE.clear()
    rag._BM25_SCORE_CACHE.clear()
